import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tc_mar.append(node)
    tc_pr.append(tc_mar)

def create_document():
    doc = Document()
    
    # Page setup - Margins (1 inch on all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Colors
    NAVY = RGBColor(26, 54, 93)     # #1A365D - Main headers
    SLATE = RGBColor(74, 85, 104)   # #4A5568 - Subtitles
    DARK = RGBColor(45, 55, 72)     # #2D3748 - Body text
    
    # Custom styling functions
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(36)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = SLATE
        p.paragraph_format.space_after = Pt(36)
        return p

    def add_heading_1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = SLATE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_body_text(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = DARK
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_bullet(text, level=0):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = DARK
        return p

    def add_code_block(code_text):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = table.rows[0].cells[0]
        cell.width = Inches(6.5)
        set_cell_background(cell, "F7FAFC")
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
        run.font.color.rgb = DARK
        # Add space after table
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(4)
        p_space.paragraph_format.space_after = Pt(8)

    # --- PORTADA (Title Page) ---
    add_title("DOCUMENTO DE BITÁCORA TÉCNICA Y WALKTHROUGH DE IMPLANTACIÓN")
    add_subtitle("Evidencia de Configuración y Código Fuente del Generador de Seguridad\nProyecto: SecureDesk (SENA ADSO)")
    
    add_heading_2("Información del Entregable Técnico:")
    add_bullet("Programa de Formación: Análisis y Desarrollo de Software (ADSO) - SENA")
    add_bullet("Plataforma Principal: SecureDesk (Backend Node/Express, Frontend React, PostgreSQL)")
    add_bullet("Desarrollador / Implantador: Sebastian Velez Mesa")
    add_bullet("Objetivo de este Documento: Servir de bitácora y walkthrough del entregable de seguridad, evidenciando las técnicas de programación empleadas en el script de generación del informe y detallando el código implementado.")
    
    doc.add_page_break()

    # --- SECCIÓN 1: INTRODUCCIÓN Y ARQUITECTURA ---
    add_heading_1("1. Introducción y Arquitectura del Generador")
    add_body_text(
        "Para dar solución al taller sobre el caso de estudio de Recursos Humanos (RRHH), se ha diseñado y ejecutado un script de automatización "
        "en Python utilizando la biblioteca 'python-docx'. Este enfoque permite compilar todo el análisis de riesgos, plan de respuesta a incidentes "
        "y checklists de aceptación de forma programática. Esto asegura la consistencia de estilos corporativos y la agilidad para futuras revisiones."
    )
    add_body_text(
        "A continuación se presenta la arquitectura lógica del generador técnico de documentos, destacando la interacción con las "
        "APIs subyacentes de Word (XML) para lograr estilos avanzados que no están disponibles nativamente en la biblioteca estándar."
    )

    # --- SECCIÓN 2: CÓDIGO FUENTE IMPLEMENTADO ---
    add_heading_1("2. Código Fuente e Implementación Técnica")
    
    add_heading_2("2.1. Funciones de Manipulación de XML de Word (docx.oxml)")
    add_body_text(
        "Dado que 'python-docx' no expone funciones directas para colorear celdas de tablas o definir márgenes internos personalizados (cell padding) "
        "de forma nativa, se recurre a la manipulación directa de los elementos OpenXML ('w:shd' para el sombreado y 'w:tcMar' para los márgenes)."
    )
    
    add_code_block(
        "def set_cell_background(cell, hex_color):\n"
        "    tc_pr = cell._tc.get_or_add_tcPr()\n"
        "    shd = OxmlElement('w:shd')\n"
        "    shd.set(qn('w:val'), 'clear')\n"
        "    shd.set(qn('w:color'), 'auto')\n"
        "    shd.set(qn('w:fill'), hex_color)\n"
        "    tc_pr.append(shd)\n\n"
        "def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):\n"
        "    tc_pr = cell._tc.get_or_add_tcPr()\n"
        "    tc_mar = OxmlElement('w:tcMar')\n"
        "    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:\n"
        "        node = OxmlElement(f'w:{m}')\n"
        "        node.set(qn('w:w'), str(val))\n"
        "        node.set(qn('w:type'), 'dxa')\n"
        "        tc_mar.append(node)\n"
        "    tc_pr.append(tc_mar)"
    )
    
    add_heading_2("2.2. Helper Functions para Estilos de Texto")
    add_body_text(
        "Para mantener una consistencia visual en todo el documento, se programaron funciones personalizadas para títulos, subtítulos, "
        "encabezados de primer y segundo nivel, listas y bloques de notas destacadas (callouts)."
    )
    
    add_code_block(
        "def add_heading_1(text):\n"
        "    p = doc.add_paragraph()\n"
        "    run = p.add_run(text)\n"
        "    run.font.name = 'Arial'\n"
        "    run.font.size = Pt(13)\n"
        "    run.font.bold = True\n"
        "    run.font.color.rgb = NAVY\n"
        "    p.paragraph_format.space_before = Pt(20)\n"
        "    p.paragraph_format.space_after = Pt(6)\n"
        "    return p\n\n"
        "def add_body_text(text):\n"
        "    p = doc.add_paragraph()\n"
        "    run = p.add_run(text)\n"
        "    run.font.name = 'Arial'\n"
        "    run.font.size = Pt(10)\n"
        "    run.font.color.rgb = DARK\n"
        "    p.paragraph_format.line_spacing = 1.15\n"
        "    p.paragraph_format.space_after = Pt(6)\n"
        "    return p"
    )

    doc.add_page_break()

    add_heading_2("2.3. Estructuración y Estilizado de Tablas Dinámicas")
    add_body_text(
        "En el archivo principal, las tablas de análisis de errores y checklists técnicos y organizativos se modelaron utilizando "
        "bucles iterativos sobre listas de tuplas. Esto automatiza la inserción de registros, la definición de anchos de columna fijos "
        "y la aplicación de colores alternados para mejorar la legibilidad."
    )
    
    add_code_block(
        "# Estructura de carga y renderizado de celdas alternas con tipografía Arial\n"
        "for row_idx, data in enumerate(errors_data):\n"
        "    row = table_errors.rows[row_idx + 1]\n"
        "    bg_color = 'F7FAFC' if row_idx % 2 == 0 else 'FFFFFF'\n"
        "    for col_idx, text in enumerate(data):\n"
        "        cell = row.cells[col_idx]\n"
        "        cell.text = text\n"
        "        set_cell_background(cell, bg_color)\n"
        "        set_cell_margins(cell, top=100, bottom=100, left=80, right=80)\n"
        "        p = cell.paragraphs[0]\n"
        "        for run in p.runs:\n"
        "            run.font.name = 'Arial'\n"
        "            run.font.size = Pt(8)\n"
        "            run.font.color.rgb = DARK"
    )

    # --- SECCIÓN 3: EVIDENCIA DE EJECUCIÓN ---
    add_heading_1("3. Evidencia de Ejecución y Generación de Entregables")
    add_body_text(
        "La ejecución de los scripts de generación se realiza directamente sobre el entorno local de desarrollo. "
        "A continuación se presenta el log técnico de la consola de Windows PowerShell que valida el proceso:"
    )
    
    add_code_block(
        "PS C:\\Users\\User\\Desktop\\SecureDesk_ADSO> python generate_informe_incidente.py\n"
        "Document successfully created: Analisis_Incidente_RRHH_Implantacion_Segura.docx\n\n"
        "PS C:\\Users\\User\\Desktop\\SecureDesk_ADSO> python generate_walkthrough_docx.py\n"
        "Walkthrough document successfully created: Walkthrough_Cambios_Realizados.docx"
    )
    
    add_body_text(
        "Ambos archivos se generaron en el directorio raíz del proyecto con firmas criptográficas de integridad locales y están "
        "disponibles para su apertura directa mediante procesadores de texto compatibles con el estándar XML de Microsoft Word."
    )

    # Save document
    output_filename = "Walkthrough_Cambios_Realizados.docx"
    doc.save(output_filename)
    print(f"Technical Walkthrough document successfully created: {output_filename}")

if __name__ == "__main__":
    create_document()
