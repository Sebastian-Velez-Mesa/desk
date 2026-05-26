import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tc_mar.append(node)
    tc_pr.append(tc_mar)

def create_document():
    doc = Document()
    
    # Page setup - Margins (1 inch on all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Colors (Modern Corporate Slate Palette)
    NAVY = RGBColor(26, 54, 93)     # #1A365D - Primary headers
    SLATE = RGBColor(74, 85, 104)   # #4A5568 - Secondary headers / subtitles
    DARK = RGBColor(45, 55, 72)     # #2D3748 - Body text
    
    # Custom styling functions
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(36)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = SLATE
        p.paragraph_format.space_after = Pt(36)
        return p

    def add_heading_1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = SLATE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_body_text(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = DARK
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_bullet(text, level=0):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = DARK
        return p

    def add_callout(text, bg_color="F7FAFC"):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = table.rows[0].cells[0]
        cell.width = Inches(6.5)
        set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        run.font.italic = True
        run.font.color.rgb = DARK
        # Add space after table
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(4)
        p_space.paragraph_format.space_after = Pt(8)

    # --- PORTADA (Title Page) ---
    add_title("INFORME TÉCNICO: ANÁLISIS DE INCIDENTE Y MARCO DE IMPLANTACIÓN SEGURA")
    add_subtitle("Caso de Estudio: Fuga de Información en Sistema de Gestión de Recursos Humanos (RRHH)\nActividad: Auditoría e Implantación Segura de Software")
    
    add_heading_2("Información del Entregable:")
    add_bullet("Programa de Formación: Análisis y Desarrollo de Software (ADSO) - SENA")
    add_bullet("Proyecto Asociado: Sistema de Información RRHH (Módulo de Talento Humano y Nómina)")
    add_bullet("Estudiante / Implantador Técnico: Sebastian Velez Mesa")
    add_bullet("Fecha de Elaboración: 26 de Mayo de 2026")
    add_bullet("Objetivo: Analizar la brecha de seguridad presentada, categorizar los fallos cometidos, definir un plan de respuesta a incidentes exhaustivo y estructurar checklists técnicos y organizativos para una aceptación segura en producción.")
    
    doc.add_page_break()

    # --- SECCIÓN 1 ---
    add_heading_1("1. Descripción General del Sistema Afectado, Contexto y Datos Comprometidos")
    
    add_heading_2("1.1. Sistema Afectado")
    add_body_text(
        "El sistema afectado corresponde a la plataforma corporativa de Gestión de Recursos Humanos (RRHH). "
        "Esta aplicación es el núcleo operativo para la administración del personal de la compañía, encargándose del registro "
        "de la información personal de los empleados, la dispersión y cálculo de la nómina mensual, el seguimiento del historial "
        "de contratos de trabajo, las evaluaciones de desempeño y la gestión de la seguridad social. Debido a sus funciones, el sistema "
        "se conecta directamente con una base de datos relacional centralizada que alberga registros confidenciales e información "
        "financiera crítica."
    )
    
    add_heading_2("1.2. Contexto de Implantación")
    add_body_text(
        "El incidente ocurrió durante la fase final de la implantación (despliegue a producción) del nuevo sistema de RRHH. "
        "Para poner en marcha la aplicación, el equipo técnico requería realizar una migración de los datos históricos de los empleados "
        "desde los servidores antiguos de la empresa hacia la base de datos de producción recién desplegada. En lugar de establecer un canal "
        "privado y seguro, y de aplicar un proceso formal de aseguramiento de la infraestructura (hardening), se optó por un método ad-hoc "
        "e inseguro para acelerar el despliegue. Asimismo, el software se puso en funcionamiento utilizando configuraciones del proveedor por defecto."
    )
    
    add_heading_2("1.3. Tipo de Datos Comprometidos")
    add_body_text(
        "La filtración expuso en un foro público de la dark web un volcado de datos completo de los empleados. La información comprometida "
        "se categoriza como Datos de Identificación Personal (PII) y Datos Financieros Sensibles bajo regulaciones de protección de datos personales:"
    )
    add_bullet("Datos de Identificación Personal (PII): Nombres completos, números de documentos de identidad (cédulas de ciudadanía), direcciones de domicilio, correos electrónicos corporativos y personales, números de contacto y contactos de emergencia.")
    add_bullet("Datos Laborales y Contractuales: Cargos, salarios asociados, historial laboral interno, fechas de vinculación, y tipo de contratación.")
    add_bullet("Datos Financieros Críticos: Números de cuentas bancarias asociadas para el depósito de nómina, certificados bancarios e información fiscal tributaria.")
    add_bullet("Datos de Seguridad Social: Afiliaciones a empresas prestadoras de salud (EPS), fondos de pensiones y cesantías, y registros de accidentes laborales.")

    # --- SECCIÓN 2 ---
    add_heading_1("2. Análisis de Errores de Seguridad y Gestión de Riesgos")
    add_body_text(
        "Durante el proceso de despliegue se cometieron múltiples fallos críticos de seguridad. A continuación se identifican y clasifican "
        "cinco de los errores más graves cometidos por el equipo técnico implantador:"
    )

    # Table 1: Errors, Risks, Analysis, Controls, Verification
    table_errors = doc.add_table(rows=6, cols=5)
    table_errors.style = 'Table Grid'
    
    headers = [
        "Error Identificado", 
        "Tipo de Riesgo Asociado", 
        "Explicación del Fallo", 
        "Control Técnico / Procedimental", 
        "Método de Verificación"
    ]
    col_widths = [Inches(1.2), Inches(1.0), Inches(1.4), Inches(1.5), Inches(1.4)]
    
    # Header styling
    hdr_cells = table_errors.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1A365D")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=80, right=80)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(255, 255, 255)

    errors_data = [
        ("Uso de credenciales de administrador por defecto del proveedor para la aplicación de RRHH.",
         "Acceso Indebido / Configuración Insegura / Trazabilidad Insuficiente",
         "Dejar las credenciales de fábrica (ej. admin/admin) permite a cualquier atacante externo o script automatizado ingresar con control total del sistema. Además, al usar todos la misma cuenta predeterminada, no es posible identificar qué persona realizó qué acción en caso de intrusión.",
         "Procedimiento de Hardening: Forzar el cambio de credenciales de administración en la primera ejecución de instalación (Setup Wizard). Implementar políticas de complejidad y MFA (Autenticación Multifactor) obligatoria.",
         "Verificación Técnica: Ejecutar un script de prueba automatizado posterior a la instalación que intente autenticarse con las claves de fábrica del fabricante, validando que el servidor responda con un código de error de login (401/403)."),
         
        ("Migración de la base de datos de empleados en texto plano mediante un servidor FTP público.",
         "Migración Insegura / Falta de Cifrado en Tránsito / Confidencialidad",
         "El protocolo FTP transmite credenciales y datos en texto plano. Utilizar una red pública de internet y un FTP sin protección expone toda la base de datos a intercepción por parte de terceros (sniffing) o indexación en buscadores de servidores FTP abiertos.",
         "Cifrado en Tránsito: Deshabilitar el protocolo FTP. Obligar al uso de SFTP (SSH File Transfer Protocol) o HTTPS (TLS 1.3) para cualquier carga de datos de la empresa, implementando listas blancas de IPs autorizadas.",
         "Verificación Técnica: Ejecutar un análisis de puertos (Nmap) en el servidor de migración para certificar que el puerto 21 está cerrado, y capturar paquetes con Wireshark en una migración de prueba para verificar que el tráfico sea ilegible."),
         
        ("Falta de cifrado de la base de datos en reposo.",
         "Falta de Cifrado en Reposo / Confidencialidad / Incumplimiento Normativo",
         "Si el archivo de base de datos se almacena en el servidor FTP o en el motor de base de datos final sin cifrar, un atacante que acceda al sistema de archivos puede leer directamente toda la PII de los empleados en texto plano.",
         "Cifrado en Reposo: Implementar Cifrado Transparente de Datos (TDE) en el motor de base de datos y cifrado a nivel de volumen del disco duro (ej. BitLocker o LUKS). Cifrar campos altamente sensibles (salarios) mediante hashes con sal.",
         "Verificación Técnica: Intentar leer los archivos binarios de la base de datos y copias de seguridad (.bak/.sql) desde el exterior y verificar mediante un visor hexadecimal que su contenido es ilegible."),
         
        ("Inexistencia de un entorno de pruebas independiente para la migración (Despliegue directo).",
         "Falta de Separación de Entornos / Configuración Insegura",
         "Realizar la migración directamente sobre el entorno productivo con datos reales de negocio, sin haber probado el canal de transmisión ni el proceso en una red aislada (Staging), incrementa significativamente la exposición a fallos lógicos e intrusiones.",
         "Separación de Entornos: Configurar entornos físicamente aislados para Desarrollo, Staging y Producción. La migración debe ensayarse primero en Staging utilizando únicamente datos ficticios u ofuscados.",
         "Verificación Técnica: Revisar el diagrama de red y la arquitectura de infraestructura del cloud para verificar que no existen rutas de red directas entre desarrollo/staging y las bases de datos de producción."),
         
        ("Inexistencia de monitorización, logs y alertas automáticas ante anomalías en la red y los servidores.",
         "Trazabilidad Insuficiente / Falta de Monitoreo",
         "La filtración tardó una semana en detectarse, y se supo por un tercero (foro de leaks). Esto demuestra que el equipo técnico no tenía alertas de descarga de grandes volúmenes de datos ni logs de accesos indebidos.",
         "Monitoreo y Logs: Configurar el registro inmutable de auditoría (logs) de accesos, inicios de sesión y descargas, y centralizarlos en un servidor syslog/SIEM, con alertas automáticas ante picos anormales de descargas.",
         "Verificación Técnica: Simular una descarga masiva de datos ficticios desde una IP no autorizada y verificar que la alerta de seguridad se active y notifique al administrador del sistema en tiempo real.")
    ]

    for row_idx, data in enumerate(errors_data):
        row = table_errors.rows[row_idx + 1]
        bg_color = "F7FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.text = text
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8)
                run.font.color.rgb = DARK
                if col_idx == 0:
                    run.font.bold = True

    # Adjust widths for Table 1
    for row in table_errors.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    # Add space after table
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(12)

    # --- SECCIÓN 3 ---
    add_heading_1("3. Plan de Respuesta a Incidentes ante la Fuga de Datos de RRHH")
    add_body_text(
        "Al detectarse que los datos de los empleados se encuentran publicados en un foro de filtraciones, la organización "
        "debe activar de inmediato el Plan de Respuesta a Incidentes (IRP) estructurado en las siguientes fases operativas:"
    )

    add_heading_2("3.1. Detección y Análisis")
    add_bullet("Validación de la Fuga: El equipo de ciberseguridad debe verificar si la muestra de datos expuesta en el foro coincide con la estructura interna y registros reales de la base de datos de empleados de la compañía.")
    add_bullet("Búsqueda Forense: Analizar los logs del servidor FTP, el servidor de aplicaciones de RRHH y el motor de base de datos PostgreSQL/MySQL para identificar la IP de origen, la marca de tiempo exacta, el método de acceso explotado y el volumen de datos sustraído.")
    add_bullet("Determinación del Alcance: Identificar exactamente cuántos empleados fueron afectados y cuáles campos específicos de información sensible fueron comprometidos (ej. si se incluyeron contraseñas de cuentas corporativas).")

    add_heading_2("3.2. Contención")
    add_bullet("Cierre Inmediato del Canal Vulnerado: Desconectar de la red y apagar el servidor FTP público. Cortar cualquier canal de transferencia de archivos expuesto a internet.")
    add_bullet("Revocación de Credenciales de Administrador: Inhabilitar la cuenta comprometida con credenciales por defecto. Forzar el cierre de todas las sesiones activas en la base de datos y la aplicación.")
    add_bullet("Aislamiento de Servidores: Aplicar reglas estrictas en el firewall del hosting (ej. bloquear todo el tráfico externo y habilitar únicamente accesos vía VPN privada para el equipo de incidentes) para evitar que el atacante intente una nueva conexión.")

    add_heading_2("3.3. Erradicación")
    add_bullet("Limpieza de Amenazas: Realizar un escaneo de malware y auditoría del sistema de archivos en busca de web shells, backdoors o scripts inyectados por el atacante para mantener persistencia.")
    add_bullet("Remediación de Vulnerabilidades: Reconfigurar la base de datos y el servidor de aplicaciones bajo guías de hardening (aseguramiento). Aplicar los controles de cambio de contraseñas de todos los usuarios de la base de datos e inhabilitar servicios redundantes.")
    add_bullet("Solicitud de Eliminación Externa: En la medida de lo posible, tomar contacto con los administradores de los repositorios de filtraciones o recurrir a servicios especializados de mitigación de marca para solicitar el retiro de la información sensible publicada.")

    add_heading_2("3.4. Recuperación")
    add_bullet("Restauración Segura de los Sistemas: Si fue necesario apagar servidores, volver a desplegar la infraestructura a partir de imágenes limpias y seguras (sin credenciales por defecto) verificadas previamente.")
    add_bullet("Validación de Integridad de Datos: Comparar el estado actual de los datos con los backups anteriores al incidente para confirmar que el atacante no modificó, inyectó o eliminó registros legítimos.")
    add_bullet("Monitoreo Continuo Post-Incidente: Establecer vigilancia extrema (24/7) sobre los logs del cortafuegos, base de datos y endpoints durante los siguientes 30 días.")
    add_bullet("Notificación Obligatoria y Comunicación Externa:")
    add_bullet("Notificar a los empleados afectados: Emitir una alerta corporativa transparente explicando lo sucedido, qué datos se comprometieron y qué medidas deben tomar (ej. cambio de contraseñas bancarias y personales).", level=1)
    add_bullet("Notificar a la autoridad competente: En concordancia con las leyes locales (como la Ley 1581 de Protección de Datos Personales en Colombia), radicar el reporte de incidente de seguridad de la información ante la entidad reguladora correspondiente dentro del término de ley.", level=1)

    add_heading_2("3.5. Lecciones Aprendidas")
    add_bullet("Sesión Post-Mortem: Reunir al equipo de desarrollo, DevOps, infraestructura y TI para analizar detalladamente la causa raíz del incidente y el tiempo de respuesta.")
    add_bullet("Actualización de Políticas y Estándares: Modificar la política de desarrollo de software para prohibir explícitamente el uso de FTP y la configuración de contraseñas por defecto.")
    add_bullet("Capacitación y Simulacros: Programar sesiones de concientización y entrenamiento técnico en seguridad en la implantación para el personal técnico.")

    doc.add_page_break()

    # --- SECCIÓN 4 ---
    add_heading_1("4. Checklist Técnico de Aceptación Segura (Antes de Producción)")
    add_body_text(
        "A continuación se define la lista de verificación técnica de aceptación segura (Checklist Técnico) que debe ejecutarse "
        "y validarse obligatoriamente de forma previa a la salida en producción de cualquier desarrollo:"
    )

    # Table 2: Technical Checklist
    table_tech = doc.add_table(rows=8, cols=5)
    table_tech.style = 'Table Grid'
    
    headers_tech = ["ID", "Control Técnico", "Descripción del Control", "Criterio de Aceptación", "Responsable"]
    col_widths_tech = [Inches(0.6), Inches(1.3), Inches(1.8), Inches(1.6), Inches(1.2)]
    
    hdr_cells_t = table_tech.rows[0].cells
    for i, title in enumerate(headers_tech):
        hdr_cells_t[i].text = title
        set_cell_background(hdr_cells_t[i], "1A365D")
        set_cell_margins(hdr_cells_t[i], top=120, bottom=120, left=80, right=80)
        p = hdr_cells_t[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(255, 255, 255)

    tech_data = [
        ("TEC-01",
         "Cambio de contraseñas por defecto",
         "Modificación obligatoria de todas las claves predeterminadas del proveedor (administración, base de datos, APIs).",
         "Cero contraseñas genéricas activas. Bloqueo de acceso al sistema si la clave es la por defecto.",
         "Administrador del Sistema / DevOps"),
         
        ("TEC-02",
         "Deshabilitación de servicios innecesarios",
         "Hardening del servidor. Apagar puertos y protocolos obsoletos o redundantes (FTP, Telnet, SSH expuesto sin llave).",
         "Escaneo de puertos nmap muestra únicamente los puertos HTTPS (443) y SSH internos (22) abiertos.",
         "DevOps / SysAdmin"),
         
        ("TEC-03",
         "Escaneo de vulnerabilidades",
         "Ejecutar escaneos automáticos de dependencias y código fuente (ej. npm audit, SonarQube, Snyk).",
         "Reporte libre de dependencias desactualizadas o exploits conocidos de severidad crítica/alta.",
         "Líder de QA / Seguridad"),
         
        ("TEC-04",
         "Remediación de fallas críticas y altas",
         "Corrección activa y aplicación de parches sobre las vulnerabilidades de seguridad detectadas en los escaneos.",
         "Validación por escáner posterior que certifique 100% de mitigación de vulnerabilidades críticas y altas.",
         "Equipo de Desarrollo"),
         
        ("TEC-05",
         "Configuración de logs centralizados",
         "Redirección de trazas de acceso, login, errores y transacciones a un servidor central de logs (SIEM / Syslog).",
         "Prueba de inyectores de logs exitosa y almacenamiento de trazas estructurado fuera del servidor de la app.",
         "SysAdmin"),
         
        ("TEC-06",
         "Cifrado de datos en tránsito (TLS)",
         "Cifrado obligatorio de todas las conexiones entrantes y salientes mediante TLS 1.2 o 1.3 con certificados SSL válidos.",
         "Redirección automática de HTTP (80) a HTTPS (443). Calificación A+ en pruebas de SSL Labs.",
         "DevOps"),
         
        ("TEC-07",
         "Separación física de entornos",
         "Configuración de redes virtuales (VPC) y servidores aislados para Desarrollo, Staging y Producción.",
         "Las credenciales y conexiones a bases de datos de producción están bloqueadas en desarrollo/staging.",
         "Arquitecto Cloud")
    ]

    for row_idx, data in enumerate(tech_data):
        row = table_tech.rows[row_idx + 1]
        bg_color = "F7FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.text = text
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8)
                run.font.color.rgb = DARK
                if col_idx == 0:
                    run.font.bold = True

    for row in table_tech.rows:
        for idx, width in enumerate(col_widths_tech):
            row.cells[idx].width = width

    # Add space after table
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(12)

    # --- SECCIÓN 5 ---
    add_heading_1("5. Checklist Organizativo de Aceptación Segura")
    add_body_text(
        "Adicionalmente a las medidas técnicas, la organización debe garantizar que se cumplan las siguientes directrices operativas "
        "y de cumplimiento legal mediante el Checklist Organizativo:"
    )

    # Table 3: Organizational Checklist
    table_org = doc.add_table(rows=7, cols=5)
    table_org.style = 'Table Grid'
    
    headers_org = ["ID", "Control Organizativo", "Descripción del Control", "Criterio de Aceptación", "Responsable"]
    col_widths_org = [Inches(0.6), Inches(1.3), Inches(1.8), Inches(1.6), Inches(1.2)]
    
    hdr_cells_o = table_org.rows[0].cells
    for i, title in enumerate(headers_org):
        hdr_cells_o[i].text = title
        set_cell_background(hdr_cells_o[i], "1A365D")
        set_cell_margins(hdr_cells_o[i], top=120, bottom=120, left=80, right=80)
        p = hdr_cells_o[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(255, 255, 255)

    org_data = [
        ("ORG-01",
         "Accesos bajo mínimo privilegio (RBAC)",
         "Garantizar que desarrolladores y operadores posean permisos mínimos requeridos para sus funciones específicas.",
         "Matriz de roles aprobada. Ningún desarrollador externo tiene permisos de escritura en la base de datos de producción.",
         "Líder de TI / Gestión de Accesos"),
         
        ("ORG-02",
         "Eliminación de cuentas temporales",
         "Borrado definitivo de credenciales o accesos creados de forma temporal para pruebas o por proveedores externos.",
         "Revisión de cuentas en bases de datos y consolas del sistema libre de perfiles de prueba o ex-empleados.",
         "Administrador de Sistemas"),
         
        ("ORG-03",
         "Documentación de configuración actualizada",
         "Mantener al día el manual de instalación, la topología de red, el inventario de software y el manual de configuración.",
         "Repositorio Git documentado y de fácil lectura con diagramas de arquitectura actualizados y aprobados.",
         "Tech Lead / DevOps"),
         
        ("ORG-04",
         "Plan de respuesta a incidentes formalizado",
         "Contar con un protocolo oficial con roles, canales de comunicación y flujos de acción claros ante brechas de datos.",
         "Documento aprobado por la gerencia y socializado con los equipos técnicos clave mediante simulacro anual.",
         "CISO / Oficial de Cumplimiento"),
         
        ("ORG-05",
         "Comunicación al responsable de datos",
         "Protocolo para notificar oportunamente incidentes de seguridad al DPO (Data Protection Officer) y entes gubernamentales.",
         "Formatos de reporte listos según la regulación de protección de datos (ej. Ley 1581). Notificaciones completadas.",
         "Oficial de Privacidad (DPO)"),
         
        ("ORG-06",
         "Prueba de restauración de backup",
         "Realizar una restauración completa de las copias de seguridad de la base de datos para comprobar que los backups sirven.",
         "Ejecutar y documentar una restauración trimestral exitosa en un entorno limpio sin pérdida de información.",
         "Administrador de Bases de Datos (DBA)")
    ]

    for row_idx, data in enumerate(org_data):
        row = table_org.rows[row_idx + 1]
        bg_color = "F7FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.text = text
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8)
                run.font.color.rgb = DARK
                if col_idx == 0:
                    run.font.bold = True

    for row in table_org.rows:
        for idx, width in enumerate(col_widths_org):
            row.cells[idx].width = width

    # Add space after table
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(12)

    # --- SECCIÓN 6 ---
    add_heading_1("6. Conclusión: Responsabilidad Profesional del Técnico Implantador")
    add_body_text(
        "La seguridad en el proceso de implantación de software no debe verse como un paso opcional, opulento o secundario "
        "en el ciclo de desarrollo de software; por el contrario, constituye una responsabilidad integral, ética y profesional del "
        "técnico implantador. El código más seguro y robusto del mundo puede ser completamente invalidado o expuesto a exploits "
        "simples si al momento del despliegue en producción se configuran incorrectamente las variables del entorno, se dejan "
        "abiertos servicios inseguros o se obvia el cambio de credenciales de fábrica del proveedor."
    )
    add_body_text(
        "El técnico implantador o especialista DevOps actúa como la última línea de defensa entre la aplicación de software en desarrollo "
        "y el entorno real expuesto al internet público. Omitir controles básicos bajo el pretexto de cumplir con plazos de tiempo ajustados "
        "o por falta de rigor organizativo es inaceptable, ya que expone a la organización a severas sanciones económicas (por "
        "violación de las leyes de protección de datos), pérdidas operativas catastróficas y al daño irreparable de la reputación institucional."
    )
    
    add_callout(
        "“La ciberseguridad en el despliegue no es una etapa adicional del proyecto, sino un atributo de calidad no funcional "
        "indispensable que garantiza la confianza y la viabilidad operativa a largo plazo de cualquier sistema de información corporativo.”"
    )

    # Save document
    output_filename = "Analisis_Incidente_RRHH_Implantacion_Segura.docx"
    doc.save(output_filename)
    print(f"Document successfully created: {output_filename}")

if __name__ == "__main__":
    create_document()
