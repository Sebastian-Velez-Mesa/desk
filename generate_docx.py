import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tc_mar.append(node)
    tc_pr.append(tc_mar)

def create_document():
    doc = Document()
    
    # Page setup - Margins (1 inch on all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Colors
    NAVY = RGBColor(26, 54, 93)     # #1A365D - Main headers
    SLATE = RGBColor(74, 85, 104)   # #4A5568 - Subtitles
    DARK = RGBColor(45, 55, 72)     # #2D3748 - Body text
    
    # Custom styling functions
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.italic = True
        run.font.color.rgb = SLATE
        p.paragraph_format.space_after = Pt(24)
        return p

    def add_heading_1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = SLATE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_body_text(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = DARK
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_bullet(text, level=0):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = DARK
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        return p

    # --- Title Page Content ---
    add_title("DOCUMENTO DE IMPLANTACIÓN SEGURA Y CONTROLES DE SEGURIDAD")
    add_subtitle("Caso de Estudio: Plataforma SecureDesk (SENA ADSO)\nActividad: Clase 4 · Controles de Seguridad para la Implantación")
    
    add_heading_2("Información de Control del Entregable:")
    add_bullet("Programa de Formación: Análisis y Desarrollo de Software (ADSO) - SENA")
    add_bullet("Proyecto Asociado: Plataforma SecureDesk (Backend Node.js/Express, Frontend React/Vite, BD PostgreSQL con Prisma ORM)")
    add_bullet("Objetivo del Documento: Definir y estructurar la descripción del sistema, sus riesgos de despliegue, el checklist de controles mínimos obligatorios y la justificación y evidencia necesarias para garantizar una implantación segura en el entorno de producción.")
    
    doc.add_page_break()

    # --- Section 1: Descripción del Sistema ---
    add_heading_1("1. Descripción General del Sistema (SecureDesk)")
    add_body_text(
        "SecureDesk es una plataforma web robusta diseñada para el comercio electrónico y la gestión operativa interna de una empresa. "
        "El sistema ha sido estructurado bajo una arquitectura desacoplada moderna: el Frontend está desarrollado en React/Vite "
        "para ofrecer una interfaz ágil, reactiva e interactiva, mientras que el Backend se ejecuta sobre Node.js y Express, gestionando "
        "las peticiones del cliente y exponiendo APIs seguras. Para la persistencia y mapeo de datos se utiliza Prisma ORM conectando "
        "con una base de datos relacional PostgreSQL. La seguridad y el resguardo de información de clientes y transacciones son los pilares "
        "fundamentales en el diseño arquitectónico de SecureDesk."
    )
    
    add_heading_2("Módulos Principales del Sistema:")
    add_bullet("Autenticación y Autorización: Implementado con tokens web de JSON (JWT) firmados digitalmente, cifrado de contraseñas con el algoritmo bcrypt en el backend, y un esquema de control de accesos basado en roles (RBAC) para diferenciar entre clientes estándar y personal administrativo.")
    add_bullet("Catálogo de Productos y Carrito de Compras: Permite la interacción ágil del cliente, búsquedas dinámicas, filtros y gestión interactiva de productos agregados antes de finalizar una compra.")
    add_bullet("Gestión de Pedidos y Transacciones (Orders): Módulo encargado de procesar la lógica de ventas, registrar órdenes de compra, validar inventario y generar un historial seguro de facturación.")
    add_bullet("Panel de Administración Técnica y Comercial: Panel protegido con permisos exclusivos que permite la gestión de inventario, edición de precios de productos y la visualización de reportes de ventas y actividad del sistema.")
    
    add_heading_2("Usuarios del Sistema:")
    add_bullet("Clientes Finales: Usuarios que acceden a la interfaz pública del frontend para explorar el catálogo y realizar compras seguras.")
    add_bullet("Administradores / Personal Comercial: Usuarios internos encargados de actualizar stock, validar pedidos e inspeccionar transacciones.")
    add_bullet("Ingenieros de Desarrollo / DevOps: Profesionales técnicos encargados del mantenimiento del código, las migraciones de base de datos y la implantación de actualizaciones.")

    add_heading_2("Contexto de Implantación:")
    add_body_text(
        "SecureDesk se despliega en una infraestructura en la nube moderna y escalable. Para garantizar la seguridad operacional, se establecen "
        "ambientes físicamente independientes de Desarrollo (local/pruebas unitarias), Staging (pruebas integrales y de regresión en un entorno similar "
        "a producción) y Producción (servidor activo con datos reales de negocio). La base de datos PostgreSQL de producción se accede mediante "
        "conexiones seguras y cifradas, aislando todas las credenciales de acceso de los repositorios de código mediante variables de entorno."
    )

    # --- Section 2: Riesgos de Seguridad Identificados ---
    add_heading_1("2. Riesgos de Seguridad en las Fases del Despliegue")
    add_body_text(
        "El proceso de implantación en producción introduce vectores de vulnerabilidad que pueden comprometer la confidencialidad, integridad "
        "y disponibilidad de la plataforma SecureDesk. A continuación se detallan los riesgos clave identificados por cada fase operativa:"
    )
    
    add_heading_2("Fase 1: Antes del Despliegue (Pre-Implantación)")
    add_bullet("Exposición de Credenciales y Secretos Criptográficos: Fuga involuntaria de variables de entorno y contraseñas críticas (ej. llave de firma JWT o credenciales de la base de datos) al subir archivos .env sin protección al control de versiones de uso compartido (GitHub).")
    add_bullet("Infiltración de Dependencias Vulnerables: Integración de paquetes npm obsoletos o maliciosos que contengan vulnerabilidades conocidas (CVE) explotables de manera remota.")
    add_bullet("Código Inestable o Inseguro: Liberación de código con malas prácticas de desarrollo (ej. falta de sanitización contra inyecciones SQL o XSS) por omisión de pruebas funcionales y de seguridad en el ciclo de desarrollo.")
    
    add_heading_2("Fase 2: Durante el Despliegue (Implantación)")
    add_bullet("Corrupción o Inconsistencia en la Base de Datos: Fallas al aplicar nuevas migraciones en producción mediante Prisma ORM, lo que podría corromper la estructura de tablas críticas (ej. orders, users) y causar pérdida parcial o total de datos transaccionales.")
    add_bullet("Degradación del Servicio o Caída (Downtime): Interrupción de la operación comercial y pérdida de disponibilidad para los clientes debido a errores imprevistos en la compilación, conflictos de configuración del hosting o falta de coordinación técnica.")
    add_bullet("Cruces e Infiltración de Datos entre Entornos: Configuración errónea que conecte el entorno de staging o desarrollo con la base de datos productiva, alterando registros reales o exponiendo información sensible de los usuarios.")
    
    add_heading_2("Fase 3: Después del Despliegue (Post-Implantación)")
    add_bullet("Ataques de Denegación de Servicio (DoS/DDoS) y Fuerza Bruta: Vulnerabilidad del servidor ante peticiones masivas en endpoints críticos (como el de autenticación /login), por ausencia de middleware de rate-limiting (limitación de tasa de peticiones).")
    add_bullet("Fugas de Información por Logging Detallado: Exposición de trazas de error completas y credenciales en archivos de log de acceso público, facilitando información interna a atacantes externos.")
    add_bullet("Incapacidad de Reacción ante Vulnerabilidades Críticas de Día Cero: Retraso en la aplicación de parches correctivos por no disponer de un canal rápido para liberar hotfixes de seguridad de emergencia.")

    doc.add_page_break()

    # --- Section 3: Tabla del Checklist de Implantación Segura ---
    add_heading_1("3. Checklist de Implantación Segura para SecureDesk")
    add_body_text(
        "A continuación se presenta la matriz de controles mínimos de seguridad obligatorios, organizada estructuradamente por categorías. "
        "Cada control incluye su clasificación temporal, su propósito técnico de protección y el incidente o riesgo directo que previene."
    )

    # Create Table: 1 Header row + 15 Data rows = 16 Rows, 5 Columns
    table = doc.add_table(rows=16, cols=5)
    table.style = 'Table Grid'
    
    headers = [
        "Categoría", 
        "Control de Seguridad", 
        "Clasificación", 
        "Para qué sirve\n(Protección / Garantía)", 
        "Qué problema previene\n(Mitigación de Incidente)"
    ]
    # Total table width should be exactly 6.5 inches (fitting 8.5" width with 1" margins)
    hdr_widths = [Inches(1.0), Inches(1.2), Inches(0.8), Inches(1.75), Inches(1.75)]
    
    # Set header content and style
    hdr_cells = table.rows[0].cells
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        set_cell_background(hdr_cells[i], "1A365D") # Dark Navy Blue
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
        
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255) # White text

    # Data rows - 15 items
    row_data = [
        # Categoría 1: Código y Ambiente
        ("Código y Ambiente", 
         "Revisión de Código y Escaneo SAST\n(ej. npm audit, SonarQube)", 
         "Preventivo",
         "Asegurar que el código fuente no posea fallos conocidos ni dependencias obsoletas con exploits públicos.",
         "Previene la intrusión de atacantes que aprovechen huecos de seguridad en dependencias externas o malas prácticas de codificación."),
        
        ("Código y Ambiente", 
         "Pruebas Funcionales y de Seguridad en Staging", 
         "Preventivo",
         "Validar el comportamiento funcional y de seguridad de la aplicación en un entorno idéntico al de producción antes de liberar.",
         "Evita fallos lógicos graves que puedan interrumpir el negocio o exponer vulnerabilidades lógicas no detectadas en producción."),
         
        ("Código y Ambiente",
         "Aislamiento y Separación Física de Ambientes",
         "Preventivo",
         "Garantizar la independencia total de las bases de datos y servidores de prueba respecto a los reales de producción.",
         "Previene la alteración accidental de datos reales de clientes por pruebas fallidas y el acceso de desarrolladores sin privilegios a producción."),
         
        ("Código y Ambiente", 
         "Variables de Entorno y Configuración de Producción Segura", 
         "Preventivo",
         "Garantizar que todas las credenciales críticas se almacenen fuera del código (en variables del host y excluidas mediante .gitignore).",
         "Evita la exposición accidental de contraseñas de bases de datos, APIs de pago o secretos de firmas JWT en repositorios públicos."),
         
        # Categoría 2: Accesos y Permisos
        ("Accesos y Permisos",
         "Control de Acceso basado en Roles (RBAC) y Seguridad API",
         "Preventivo",
         "Restringir las rutas y recursos del servidor únicamente a usuarios con roles explícitamente autorizados (ej. Administradores).",
         "Evita escalada de privilegios y accesos no autorizados a operaciones sensibles, como manipulación del inventario o alteración de pedidos."),
         
        ("Accesos y Permisos",
         "Diferenciación y Rotación Estricta de Credenciales",
         "Preventivo",
         "Garantizar el uso de credenciales únicas para producción, impidiendo el uso de claves débiles o contraseñas usadas en desarrollo.",
         "Previene que un compromiso del entorno local de desarrollo afecte o permita el control no autorizado de la base de datos de producción."),
         
        ("Accesos y Permisos",
         "Autenticación de Doble Factor (MFA) en Cuentas Administrativas",
         "Preventivo",
         "Exigir un segundo factor de autenticación dinámico (OTP/Authenticator) para ingresar a consolas cloud y paneles de administración.",
         "Previene el acceso ilícito a la infraestructura por robo de credenciales de administradores mediante phishing o fuerza bruta."),

        # Categoría 3: Datos y Respaldos
        ("Datos y Respaldos",
         "Cifrado de Datos Sensibles (bcrypt y SSL/TLS en Base de Datos)",
         "Preventivo",
         "Garantizar la confidencialidad de la información crítica (como contraseñas) mediante hash robusto e impedir la lectura en tránsito.",
         "Previene el robo de contraseñas en texto plano en caso de brechas en la base de datos y ataques tipo Man-in-the-Middle (MitM)."),
         
        ("Datos y Respaldos",
         "Backup de Seguridad Previo al Despliegue",
         "Preventivo",
         "Garantizar que exista una copia exacta de la base de datos de producción antes de realizar cambios de código o esquema de Prisma.",
         "Evita la pérdida catastrófica de información o la inoperabilidad permanente del sistema por fallos en las migraciones de base de datos."),
         
        ("Datos y Respaldos",
         "Plan de Rollback y Restauración Automatizada de Backups",
         "Correctivo",
         "Restaurar la aplicación al último estado funcional conocido en caso de que ocurra una falla crítica post-despliegue.",
         "Mitiga el tiempo prolongado de inactividad (downtime), restableciendo la disponibilidad operativa rápidamente ante fallas no previstas."),
         
        # Categoría 4: Aprobaciones y Comunicación
        ("Aprobaciones y Comunicación",
         "Aprobación Escrita y Formal del Responsable Técnico",
         "Preventivo",
         "Formalizar la revisión del cumplimiento de todos los pre-requisitos antes de proceder al despliegue en producción.",
         "Previene despliegues apresurados, desorganizados o no autorizados que puedan desestabilizar la operación comercial."),
         
        ("Aprobaciones y Comunicación",
         "Definición de Ventana de Mantenimiento Programada",
         "Preventivo",
         "Establecer un horario de bajo tráfico transaccional (ej. madrugadas) para realizar la actualización de la plataforma.",
         "Previene pérdidas económicas mayores por caídas de servicio durante horas pico de ventas y transacciones de clientes."),
         
        ("Aprobaciones y Comunicación",
         "Notificación y Comunicación Previa a Usuarios y Clientes",
         "Preventivo",
         "Informar con antelación a los clientes y personal operativo sobre los cortes breves programados del servicio.",
         "Evita la insatisfacción y pérdida de confianza de los clientes, así como la saturación de los canales de soporte técnico."),

        # Categoría 5: Monitoreo y Continuidad
        ("Monitoreo y Continuidad",
         "Monitoreo Activo, Logging y Alertas Post-Despliegue",
         "Detectivo",
         "Registrar eventos, accesos y errores del backend e implementar alertas automáticas ante anomalías o picos de error 5xx.",
         "Permite identificar intrusiones, intentos de inyección de código o fallos del backend de manera inmediata antes de que causen daños mayores."),

        ("Monitoreo y Continuidad",
         "Plan de Parches de Emergencia y Continuidad del Negocio",
         "Correctivo",
         "Establecer una ruta rápida (hotfix) para desplegar correcciones críticas de seguridad omitiendo el ciclo estándar.",
         "Previene la exposición prolongada ante fallas graves del día cero (Zero-Day) o fallos de infraestructura del proveedor hosting.")
    ]

    # Style and populate table rows
    for row_idx, data in enumerate(row_data):
        row = table.rows[row_idx + 1]
        
        # Alternating background colors for rows
        bg_color = "F7FAFC" if row_idx % 2 == 0 else "FFFFFF"
        
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.text = text
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
            
            # Formatting text inside the cells
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = DARK
                    
                    # Apply specific style to Classification column (col 2)
                    if col_idx == 2:
                        run.font.bold = True
                        if text == "Preventivo":
                            run.font.color.rgb = RGBColor(43, 108, 176) # Slate Blue
                        elif text == "Detectivo":
                            run.font.color.rgb = RGBColor(214, 158, 46) # Warm Gold/Orange
                        elif text == "Correctivo":
                            run.font.color.rgb = RGBColor(229, 62, 62)   # Soft Red

    # Apply widths to all cells in the table
    for row in table.rows:
        for idx, width in enumerate(hdr_widths):
            row.cells[idx].width = width

    # Add space after the table
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(12)

    # --- Section 4: Justificación y Explicación Escrita ---
    add_heading_1("4. Análisis Detallado y Justificación de los Controles")
    add_body_text(
        "En concordancia con los estándares académicos y profesionales requeridos, la matriz presentada anteriormente no "
        "consiste en un conjunto aislado de verificaciones simples, sino en un marco integrado de defensa en profundidad. "
        "A continuación se detalla la justificación operativa, arquitectónica y metodológica de cada categoría y clasificación "
        "de controles implementada:"
    )

    add_heading_2("4.1. Controles Preventivos: Reducción de la Superficie de Ataque y Mitigación Proactiva")
    add_body_text(
        "Los controles clasificados como Preventivos son aquellos diseñados para evitar activamente que un riesgo llegue a "
        "materializarse en el entorno operativo. En SecureDesk, el grueso de las defensas es preventivo, fundamentado bajo "
        "el principio de 'diseño seguro desde el origen'."
    )
    add_body_text(
        "Al separar de forma física los entornos (Desarrollo, Staging, Producción), garantizamos que las pruebas experimentales "
        "o la inestabilidad de una funcionalidad nueva jamás toquen o expongan los datos reales de los clientes del e-commerce. "
        "Esto se complementa con la inyección segura de variables de entorno y el uso estricto del archivo '.gitignore', evitando "
        "el riesgo sumamente común de fuga de credenciales críticas en repositorios compartidos."
    )
    add_body_text(
        "En la capa del software, la implementación de cifrado robusto para contraseñas usando el algoritmo 'bcrypt' y de comunicación "
        "segura HTTPS/SSL previene que una eventual brecha de la base de datos comprometa la identidad de los usuarios, o que la información "
        "del carrito e historial de compras sea interceptada por atacantes de red en el tránsito hacia el backend. Asimismo, los controles "
        "de acceso lógico basados en roles (RBAC) y la autenticación multifactor (MFA) blindan el panel de administración, evitando "
        "que un atacante realice escalada de privilegios o se apodere de la consola en la nube."
    )

    add_heading_2("4.2. Control Detectivo: Monitoreo Continuo e Identificación Temprana")
    add_body_text(
        "Dado que ninguna estrategia preventiva es infalible ante ataques dirigidos o fallas imprevistas, se incluye el control de "
        "'Monitoreo Activo, Logging y Alertas Post-Despliegue' con clasificación de tipo Detectivo. Este control tiene como propósito "
        "visibilizar el comportamiento interno de la aplicación en tiempo real."
    )
    add_body_text(
        "Mediante el registro estructurado de eventos de acceso y la tasa de códigos de error HTTP del backend, el equipo técnico es capaz "
        "de identificar patrones anómalos (por ejemplo, cientos de peticiones fallidas consecutivas al endpoint de login, indicando un "
        "ataque de fuerza bruta en curso, o una súbita oleada de errores 500 post-despliegue que denota una falla de código). "
        "La detección temprana permite reaccionar velozmente, minimizando el impacto temporal antes de que la seguridad se vea gravemente "
        "comprometida."
    )

    add_heading_2("4.3. Controles Correctivos: Resiliencia, Recuperación y Continuidad Operativa")
    add_body_text(
        "Los controles de tipo Correctivo se activan únicamente una vez que el incidente se ha manifestado en el entorno productivo. "
        "Su objetivo primordial es mitigar el daño y restablecer el estado seguro y operativo normal del sistema lo antes posible."
    )
    add_body_text(
        "En SecureDesk, esto se materializa mediante dos pilares: el Plan de Rollback Automatizado (asociado a la restauración del backup "
        "en frío de PostgreSQL y la reversión de imágenes de Docker del backend) y el Plan de Parches de Emergencia. Si una migración del "
        "ORM Prisma corrompe registros de transacciones o el nuevo código introduce una vulnerabilidad crítica explotable en producción, "
        "los controles correctivos facultan al equipo DevOps para devolver la plataforma a la versión estable previa en cuestión de minutos, "
        "o liberar un hotfix inmediato, minimizando drásticamente las pérdidas económicas y el tiempo de inactividad comercial (downtime)."
    )

    doc.add_page_break()

    # --- Section 5: Importancia de la Evidencia ---
    add_heading_1("5. La Importancia de la Evidencia en el Control y la Implantación")
    add_body_text(
        "Un principio rector del aseguramiento de la información en el desarrollo de software establece que: 'un control sin "
        "evidencia no sostiene la implantación'. Esto significa que no basta con planificar e implementar teóricamente las medidas de "
        "seguridad; es imperativo documentar y conservar pruebas tangibles que certifiquen su correcta ejecución."
    )
    
    add_heading_2("Razones por las cuales es obligatorio documentar las evidencias:")
    add_bullet(
        "Demostración de Cumplimiento Técnico (Auditoría): En entornos empresariales reales, las plataformas de software "
        "deben someterse a auditorías internas y externas para cumplir con regulaciones (como ISO 27001 o la Ley de Protección de Datos). "
        "Las evidencias son las pruebas irrefutables presentadas ante los auditores."
    )
    add_bullet(
        "Trazabilidad e Investigación Forense: Si ocurre un incidente de seguridad en el futuro, las marcas de tiempo y los reportes "
        "documentados del despliegue permiten rastrear el estado exacto de las configuraciones y determinar si el fallo se debió a un "
        "cambio reciente o a una omisión en el checklist."
    )
    add_bullet(
        "Reducción del Factor de Error Humano: El deber de adjuntar evidencias escritas y capturas de consola fuerza al equipo "
        "de desarrollo a realizar cada verificación de manera minuciosa, desalentando la omisión de pasos por prisa o exceso de confianza."
    )
    add_bullet(
        "Transparencia y Responsabilidad Compartida: Las aprobaciones formales escritas del responsable técnico evitan que el despliegue "
        "sea una acción aislada e inconsulta de un desarrollador. Establecen una corresponsabilidad institucional en la liberación del software."
    )

    add_heading_2("Evidencias Específicas Requeridas para SecureDesk:")
    add_bullet("Evidencia del Análisis SAST: Reporte en formato PDF generado por SonarQube o captura de la consola con la salida de 'npm audit' mostrando cero vulnerabilidades críticas antes del despliegue.")
    add_bullet("Evidencia de Cifrado y Configuración: Captura de pantalla de la configuración del entorno de staging o producción (con datos confidenciales ofuscados) donde se aprecie la inicialización de variables de entorno seguras (DATABASE_URL con SSL activo y JWT_SECRET de alta entropía).")
    add_bullet("Evidencia del Respaldo (Backup): Captura de pantalla del log de PostgreSQL confirmando la creación exitosa del archivo de volcado (.sql o backup binario) con su respectiva huella de tiempo (timestamp) previa a la ventana de mantenimiento.")
    add_bullet("Evidencia de Aprobación Formal: Ticket de soporte (Jira/Trello) aprobado por el Tech Lead o correo institucional donde se autorice explícitamente el despliegue adjuntando los resultados satisfactorios de staging.")

    # Save document
    output_filename = "Checklist_Implantacion_Segura_SecureDesk.docx"
    doc.save(output_filename)
    print(f"Document successfully created: {output_filename}")

if __name__ == "__main__":
    create_document()
